#!/usr/bin/env python3
"""Extract paragraphs and table structure from a DOCX file."""

from __future__ import annotations

import argparse
import json
from pathlib import Path

from docx import Document


def extract(path: Path) -> dict:
    document = Document(path)
    paragraphs = []
    for index, paragraph in enumerate(document.paragraphs):
        text = paragraph.text.strip()
        if text:
            paragraphs.append(
                {
                    "index": index,
                    "text": text,
                    "style": getattr(paragraph.style, "name", ""),
                }
            )

    tables = []
    for table_index, table in enumerate(document.tables):
        rows = []
        for row_index, row in enumerate(table.rows):
            cells = []
            for cell_index, cell in enumerate(row.cells):
                cells.append(
                    {
                        "cell_index": cell_index,
                        "text": cell.text.strip(),
                    }
                )
            rows.append({"row_index": row_index, "cells": cells})
        tables.append({"table_index": table_index, "rows": rows})

    return {"path": str(path), "paragraphs": paragraphs, "tables": tables}


def main() -> int:
    parser = argparse.ArgumentParser(description="Extract DOCX paragraphs and tables as JSON.")
    parser.add_argument("docx_path", help="Path to a .docx file")
    parser.add_argument("--pretty", action="store_true", help="Pretty-print JSON")
    args = parser.parse_args()

    path = Path(args.docx_path).resolve()
    if not path.exists():
        raise SystemExit(f"File not found: {path}")

    data = extract(path)
    if args.pretty:
        print(json.dumps(data, ensure_ascii=False, indent=2))
    else:
        print(json.dumps(data, ensure_ascii=False))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
