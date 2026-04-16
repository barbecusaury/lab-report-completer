#!/usr/bin/env python3
"""Identify likely fill targets in a DOCX report template."""

from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document

TARGET_KEYWORDS = (
    "实验目的",
    "实验内容",
    "实验设备",
    "软件环境",
    "实验过程",
    "结果",
    "异常问题",
    "解决方案",
    "实验总结",
    "截图",
    "姓名",
    "学号",
    "班级",
    "专业",
)

PLACEHOLDER_TOKENS = (
    "请写",
    "待补",
    "模板",
    "说明",
    "红字",
    "可贴图",
    "年  月  日",
)


def main() -> int:
    parser = argparse.ArgumentParser(description="Inspect likely fill targets in a DOCX report.")
    parser.add_argument("docx_path", help="Path to a .docx file")
    args = parser.parse_args()

    path = Path(args.docx_path).resolve()
    if not path.exists():
        raise SystemExit(f"File not found: {path}")

    document = Document(path)

    print(f"DOCX: {path}")
    print("Paragraph targets:")
    found_any = False
    for index, paragraph in enumerate(document.paragraphs):
        text = paragraph.text.strip()
        if not text:
            continue
        if any(keyword in text for keyword in TARGET_KEYWORDS) or any(
            token in text for token in PLACEHOLDER_TOKENS
        ):
            found_any = True
            print(f"  P{index}: {text}")

    print("Table cell targets:")
    for table_index, table in enumerate(document.tables):
        for row_index, row in enumerate(table.rows):
            for cell_index, cell in enumerate(row.cells):
                text = cell.text.strip()
                if not text:
                    continue
                if any(keyword in text for keyword in TARGET_KEYWORDS) or any(
                    token in text for token in PLACEHOLDER_TOKENS
                ):
                    found_any = True
                    print(f"  T{table_index} R{row_index} C{cell_index}: {text}")

    if not found_any:
        print("  No obvious fill targets detected.")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
