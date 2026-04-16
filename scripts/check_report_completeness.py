#!/usr/bin/env python3
"""Perform lightweight completeness checks for a report DOCX."""

from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document

SUSPICIOUS_PHRASES = (
    "请写出",
    "请写",
    "本模板",
    "红色字部分",
    "红字",
    "其他要求",
    "TODO",
)

SECTION_KEYWORDS = (
    "实验目的",
    "实验内容",
    "实验设备",
    "软件环境",
    "实验过程",
    "操作异常问题",
    "实验总结",
)


def collect_lines(document: Document) -> list[str]:
    lines: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            lines.append(text)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    lines.append(text)
    return lines


def main() -> int:
    parser = argparse.ArgumentParser(description="Check for obvious report completeness issues.")
    parser.add_argument("docx_path", help="Path to a .docx file")
    args = parser.parse_args()

    path = Path(args.docx_path).resolve()
    if not path.exists():
        raise SystemExit(f"File not found: {path}")

    document = Document(path)
    lines = collect_lines(document)
    issues: list[str] = []

    for phrase in SUSPICIOUS_PHRASES:
        if any(phrase in line for line in lines):
            issues.append(f"Suspicious leftover phrase found: {phrase}")

    for keyword in SECTION_KEYWORDS:
        if not any(keyword in line for line in lines):
            issues.append(f"Section keyword not found: {keyword}")

    if issues:
        print("Issues found:")
        for issue in issues:
            print(f"- {issue}")
        return 1

    print("No obvious completeness issues found.")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
